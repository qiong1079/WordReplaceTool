import os
import re
import sys
from docx import Document
from win32com.client import Dispatch  # 需要安装pywin32

# 配置替换规则
REPLACE_RULES = {
    "需要更换的字段": "更换为的字段",
    # 添加更多规则...
}

SOURCE_FOLDER = r"资源路径"
OUTPUT_FOLDER = r"输出路径"
LOG_FILE = os.path.join(OUTPUT_FOLDER, "processing.log")  # 日志文件路径

class ReplacementCounter:
    """替换次数计数器"""
    def __init__(self):
        self.count = 0
        self.details = {}

    def add(self, old_word, times=1):
        self.count += times
        self.details[old_word] = self.details.get(old_word, 0) + times

def convert_doc_to_docx(input_path, output_path):
    """将doc文件转换为docx格式"""
    try:
        word = Dispatch('Word.Application')
        doc = word.Documents.Open(input_path)
        doc.SaveAs(output_path, FileFormat=16)  # 16表示docx格式
        doc.Close()
        word.Quit()
        return True
    except Exception as e:
        print(f"DOC转换失败: {str(e)}")
        return False

def smart_replace(text, rules, counter):
    """执行智能替换并统计次数"""
    for old, new in rules.items():
        replaced_text, replacements = re.subn(
            re.escape(old), 
            new, 
            text, 
            flags=re.IGNORECASE
        )
        if replacements > 0:
            text = replaced_text
            counter.add(old, replacements)
    return text

def process_element(element, rules, counter):
    """处理文档元素"""
    if hasattr(element, 'runs'):
        for run in element.runs:
            original = run.text
            if not original.strip():
                continue
            modified = smart_replace(original, rules, counter)
            if modified != original:
                run.text = modified

def process_document(doc_path, counter):
    """处理单个文档"""
    doc = Document(doc_path)
    
    # 处理文档属性
    core_props = doc.core_properties
    for prop in ['title', 'subject', 'keywords']:
        value = getattr(core_props, prop, '')
        if value:
            setattr(core_props, prop, smart_replace(value, REPLACE_RULES, counter))
    
    # 处理正文
    for para in doc.paragraphs:
        process_element(para, REPLACE_RULES, counter)
    
    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_element(para, REPLACE_RULES, counter)
    
    # 处理页眉页脚
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header:
                for para in header.paragraphs:
                    process_element(para, REPLACE_RULES, counter)
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    process_element(para, REPLACE_RULES, counter)
    
    return doc

def batch_process():
    """增强版批量处理"""
    # 初始化环境
    if not os.path.exists(OUTPUT_FOLDER):
        os.makedirs(OUTPUT_FOLDER)
    
    total_files = 0
    success_files = 0
    error_log = []
    temp_files = []

    # 处理日志文件
    log_handle = open(LOG_FILE, 'w', encoding='utf-8') if LOG_FILE else None

    print("┌──────────────────────────────┬────────┬───────┐")
    print("│          文件名             │ 替换数 │ 状态  │")
    print("├──────────────────────────────┼────────┼───────┤")

    # 遍历源文件夹
    for filename in os.listdir(SOURCE_FOLDER):
        base, ext = os.path.splitext(filename)
        if ext.lower() not in ('.doc', '.docx'):
            continue

        total_files += 1
        input_path = os.path.join(SOURCE_FOLDER, filename)
        output_name = f"{base}.docx" if ext == '.doc' else filename
        output_path = os.path.join(OUTPUT_FOLDER, output_name)
        counter = ReplacementCounter()
        doc = None

        try:
            # 处理DOC文件
            if ext.lower() == '.doc':
                temp_path = os.path.join(OUTPUT_FOLDER, f"temp_{base}.docx")
                if convert_doc_to_docx(input_path, temp_path):
                    doc = process_document(temp_path, counter)
                    temp_files.append(temp_path)
                else:
                    raise Exception("DOC转换失败")
            else:
                doc = process_document(input_path, counter)
            
            # 保存文档
            if doc:
                doc.save(output_path)
                status = "成功"
                success_files += 1
            else:
                status = "失败"
            
            # 输出结果
            log = f"{filename.ljust(30)} | 替换{counter.count}次 | {status}"
            print(f"│ {filename[:28].ljust(28)} │ {str(counter.count).center(6)} │ \033[32m{status}\033[0m │")
            
        except Exception as e:
            error_type = type(e).__name__
            error_msg = f"{error_type}: {str(e)}"
            error_log.append(f"{filename}: {error_msg}")
            status = "失败"
            print(f"│ {filename[:28].ljust(28)} │ {'--'.center(6)} │ \033[31m{status}\033[0m │")
            if log_handle:
                log_handle.write(f"[ERROR] {filename}: {error_msg}\n")

    # 清理临时文件
    for temp in temp_files:
        try:
            os.remove(temp)
        except:
            pass

    # 关闭日志文件
    if log_handle:
        log_handle.close()

    # 最终统计
    print("└──────────────────────────────┴────────┴───────┘")
    print(f"\n处理完成：共处理 {total_files} 个文件")
    print(f"成功：{success_files} 个")
    print(f"失败：{total_files - success_files} 个")
    
    if error_log:
        print("\n错误详情：")
        for err in error_log:
            print(f" • {err}")

    if LOG_FILE:
        print(f"\n详细日志已保存至：{LOG_FILE}")

if __name__ == "__main__":
    print("=== 专业文档处理工具 ===")
    print("支持格式：DOC, DOCX")
    print("当前替换规则：")
    for i, (old, new) in enumerate(REPLACE_RULES.items(), 1):
        print(f"{i}. 将「{old}」替换为「{new}」")
    
    batch_process()
    print("=== 操作结束 ===")
